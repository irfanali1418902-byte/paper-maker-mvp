"""Exam paper export service (Word + PDF).

build_paper_docx assembles a printable .docx with python-docx, mirroring the
layout of static/print.html (letterhead, meta block, numbered bilingual
questions). build_paper_pdf renders that same .docx to PDF via a headless
LibreOffice subprocess — one layout, two formats.

Urdu support: Urdu runs get the complex-script font set (w:rFonts w:cs) plus
a w:rtl flag, and their paragraphs are marked w:bidi so Word lays them out
right-to-left. The target Nastaliq font must also be installed on the box
that *renders* the PDF (LibreOffice substitutes a fallback otherwise).
"""

import base64
import io
import json
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.repositories import papers_repository, questions_repository, settings_repository
from app.services.exceptions import PdfConversionFailed

# Pakistani schools mein de-facto Nastaliq font. Viewer/LibreOffice ke paas
# na ho to woh apna fallback substitute kar deta hai.
URDU_FONT = "Jameel Noori Nastaleeq"
DEFAULT_ACCENT = "0E4D3C"
# LibreOffice ek fresh profile har baar ~50s leta hai; ek persistent profile
# reuse karne se warm conversions ~13s par aa jaati hain. Timeout cold first
# run (profile banti hai) ko cover karne ke liye generous rakha hai.
PDF_TIMEOUT_SECONDS = 120
_LO_PROFILE_DIR = Path(tempfile.gettempdir()) / "paper_maker_libreoffice_profile"


# ---- public API -------------------------------------------------------------


def build_paper_docx(paper_id: str) -> Optional[bytes]:
    """Returns the paper as .docx bytes, or None if no paper with that id
    (route maps to 404)."""
    context = _load_paper(paper_id)
    if context is None:
        return None
    paper, questions, settings = context

    doc = Document()
    _render_paper(doc, paper, questions, settings)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_paper_pdf(paper_id: str) -> Optional[bytes]:
    """Builds the .docx then converts it to PDF via LibreOffice. None if the
    paper doesn't exist (404); raises PdfConversionFailed if the converter
    isn't available or the conversion errors (route maps to 502)."""
    docx_bytes = build_paper_docx(paper_id)
    if docx_bytes is None:
        return None
    return _convert_docx_to_pdf(docx_bytes)


# ---- data loading -----------------------------------------------------------


def _load_paper(paper_id: str) -> Optional[tuple[dict, list[dict], dict]]:
    paper = papers_repository.find_by_id(paper_id)
    if paper is None:
        return None
    question_ids = json.loads(paper["question_ids"])
    questions: list[dict] = []
    for qid in question_ids:
        q = questions_repository.find_by_id(qid)
        if q is not None:
            questions.append(q)
    settings = settings_repository.get() or {}
    return paper, questions, settings


# ---- docx rendering ---------------------------------------------------------


def _render_paper(doc, paper: dict, questions: list[dict], settings: dict) -> None:
    accent = _clean_hex(settings.get("accent_color")) or DEFAULT_ACCENT

    _render_letterhead(doc, settings, accent)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{paper['subject']} — Examination Paper")
    run.bold = True
    run.underline = True
    run.font.size = Pt(15)

    _render_meta(doc, paper)

    instr = doc.add_paragraph()
    _style_run(instr.add_run("Attempt all questions. "), size=10, italic=True)
    _style_run(instr.add_run("تمام سوالات کے جواب دیں۔"), size=11, urdu=True)

    for i, q in enumerate(questions, start=1):
        _render_question(doc, i, q)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(footer.add_run("Generated with AII Smart Paper Maker"), size=8, color="888888")


def _render_letterhead(doc, settings: dict, accent: str) -> None:
    logo = _decode_logo(settings.get("logo_base64"))
    if logo is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(io.BytesIO(logo), width=Inches(0.9))
        except Exception:
            # Corrupt/unsupported image data shouldn't sink the whole export.
            pass

    name_en = settings.get("school_name") or ""
    if name_en:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(p.add_run(name_en), size=20, bold=True, color=accent)

    name_ur = settings.get("school_name_ur") or ""
    if name_ur:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _make_para_rtl(p)
        _style_run(p.add_run(name_ur), size=17, bold=True, urdu=True)

    address = settings.get("address") or ""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if address:
        _style_run(p.add_run(address), size=9, color="555555")
    _add_bottom_border(p, color=accent, sz="18")


def _render_meta(doc, paper: dict) -> None:
    blank = "____________"
    rows = [
        [
            f"Name / نام: {blank}",
            f"Class / جماعت: {paper.get('class_name') or blank}",
            f"Date / تاریخ: {blank}",
        ],
        [
            f"Roll No / رول نمبر: {blank}",
            f"Total Marks / کل نمبر: {paper['total_marks']}",
            f"Time / وقت: {blank}",
        ],
    ]
    table = doc.add_table(rows=2, cols=3)
    for r, cells in enumerate(rows):
        for c, text in enumerate(cells):
            cell = table.cell(r, c)
            _style_run(cell.paragraphs[0].add_run(text), size=10)


def _render_question(doc, number: int, q: dict) -> None:
    head = doc.add_paragraph()
    head.paragraph_format.space_before = Pt(8)
    head.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    _style_run(head.add_run(f"Q{number}. "), size=12, bold=True)
    _style_run(head.add_run(q.get("question_en") or ""), size=12)
    _style_run(head.add_run(f"\t[{q['marks']}]"), size=10, color="555555")

    if q.get("question_ur"):
        p = doc.add_paragraph()
        _make_para_rtl(p)
        _style_run(p.add_run(q["question_ur"]), size=14, urdu=True)

    if q.get("visual_emoji") and q.get("visual_count"):
        p = doc.add_paragraph()
        _style_run(p.add_run(q["visual_emoji"] * int(q["visual_count"])), size=20)

    options_en = _safe_json_list(q.get("options_en"))
    options_ur = _safe_json_list(q.get("options_ur"))
    if options_en:
        _render_options(doc, options_en, options_ur)
    else:
        space = doc.add_paragraph()
        _add_bottom_border(space, color="999999", sz="6", val="dotted")


def _render_options(doc, options_en: list, options_ur: list) -> None:
    # Two-column grid, filled left-to-right — matches print.html's layout.
    row_count = (len(options_en) + 1) // 2
    table = doc.add_table(rows=row_count, cols=2)
    for idx, opt_en in enumerate(options_en):
        cell = table.cell(idx // 2, idx % 2)
        text = f"◯  {opt_en}"
        if idx < len(options_ur) and options_ur[idx]:
            text += f" / {options_ur[idx]}"
        _style_run(cell.paragraphs[0].add_run(text), size=11)


# ---- python-docx run/paragraph helpers --------------------------------------


def _style_run(run, *, size=None, bold=False, italic=False, color=None, urdu=False):
    """Single entry point for run formatting. urdu=True wires up the
    complex-script font + right-to-left flag so Nastaliq lays out correctly."""
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)

    if urdu:
        run.font.name = URDU_FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), URDU_FONT)
        rfonts.set(qn("w:hAnsi"), URDU_FONT)
        rfonts.set(qn("w:cs"), URDU_FONT)  # complex-script font = the Urdu one
        rpr.append(OxmlElement("w:rtl"))
        rpr.append(OxmlElement("w:cs"))
    return run


def _make_para_rtl(paragraph) -> None:
    """Mark a paragraph as bidi (right-to-left) and right-align it."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def _add_bottom_border(paragraph, *, color: str, sz: str, val: str = "single") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), val)
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


# ---- small pure helpers -----------------------------------------------------


def _clean_hex(value: Optional[str]) -> Optional[str]:
    """'#0e4d3c' -> '0E4D3C' (RGBColor.from_string wants no hash). None/blank
    -> None so the caller can fall back to the default accent."""
    if not value:
        return None
    cleaned = value.strip().lstrip("#").upper()
    return cleaned if len(cleaned) == 6 else None


def _safe_json_list(raw: Optional[str]) -> list:
    if not raw:
        return []
    try:
        parsed = json.loads(raw)
    except (json.JSONDecodeError, TypeError):
        return []
    return parsed if isinstance(parsed, list) else []


def _decode_logo(logo_base64: Optional[str]) -> Optional[bytes]:
    """Accepts a raw base64 string or a data URL ('data:image/png;base64,...').
    Returns the decoded bytes, or None if absent/undecodable."""
    if not logo_base64:
        return None
    payload = logo_base64.split(",", 1)[1] if "," in logo_base64 else logo_base64
    try:
        return base64.b64decode(payload)
    except (ValueError, TypeError):
        return None


# ---- PDF conversion (LibreOffice headless) ----------------------------------


def _find_soffice() -> Optional[str]:
    """Locate the LibreOffice binary: SOFFICE_PATH env override, then PATH,
    then the usual install locations on Windows/Linux."""
    override = os.environ.get("SOFFICE_PATH")
    if override and Path(override).exists():
        return override

    found = shutil.which("soffice") or shutil.which("soffice.exe")
    if found:
        return found

    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/opt/libreoffice/program/soffice",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    soffice = _find_soffice()
    if soffice is None:
        raise PdfConversionFailed(
            "LibreOffice (soffice) nahi mila. PDF export ke liye server par "
            "LibreOffice install karein ya SOFFICE_PATH env set karein. "
            "(Word export bina LibreOffice ke chalta hai.)"
        )

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        docx_file = tmp_path / "paper.docx"
        docx_file.write_bytes(docx_bytes)

        # Persistent profile (system temp mein fixed path) — har call par dobara
        # banane ki bajaye reuse hota hai, isse warm conversions tez (~13s) ho
        # jaati hain. Assumption: ek waqt mein ek hi conversion (single-school
        # MVP); LibreOffice profile ko lock karta hai, concurrent calls clash
        # kar sakti hain — async out of scope hai (ARCHITECTURE.md §8).
        _LO_PROFILE_DIR.mkdir(parents=True, exist_ok=True)
        profile_uri = _LO_PROFILE_DIR.as_uri()
        cmd = [
            soffice,
            "--headless",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(docx_file),
        ]
        try:
            proc = subprocess.run(cmd, capture_output=True, timeout=PDF_TIMEOUT_SECONDS)
        except subprocess.TimeoutExpired as e:
            raise PdfConversionFailed(
                f"PDF conversion {PDF_TIMEOUT_SECONDS}s mein finish nahi hui (LibreOffice timeout)."
            ) from e

        pdf_file = tmp_path / "paper.pdf"
        if proc.returncode != 0 or not pdf_file.exists():
            detail = proc.stderr.decode(errors="ignore").strip()[:300] or "unknown error"
            raise PdfConversionFailed(f"LibreOffice conversion fail hui: {detail}")
        return pdf_file.read_bytes()
