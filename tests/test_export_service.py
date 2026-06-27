"""Tests for export_service — .docx structure/Urdu wiring + PDF error paths.

The actual PDF rendering needs LibreOffice installed, which CI/dev machines
may lack; we test the conversion's contract (404 passthrough, missing-binary
error) by stubbing the binary lookup rather than shelling out.
"""

import io

import pytest
from docx import Document

from app.repositories import papers_repository, questions_repository, settings_repository
from app.services import export_service
from app.services.exceptions import PdfConversionFailed


def _insert_question(qid: str, marks: int, **overrides):
    row = {
        "id": qid,
        "subject": "Science",
        "topic": "Plants",
        "bloom_level": "REMEMBER",
        "difficulty": "easy",
        "question_type": "multiple-choice",
        "marks": marks,
        "question_en": "What is photosynthesis?",
        "question_ur": "ضوئی تالیف کیا ہے؟",
        "options_en": '["Light", "Water", "Air", "Soil"]',
        "options_ur": '["روشنی", "پانی", "ہوا", "مٹی"]',
        "correct_answer_en": "Light",
        "correct_answer_ur": "روشنی",
        "explanation_en": None,
        "explanation_ur": None,
        "visual_emoji": None,
        "visual_count": None,
    }
    row.update(overrides)
    questions_repository.insert(row)


def _setup_paper(test_db):
    _insert_question("q1", marks=2)
    _insert_question(
        "q2",
        marks=5,
        question_en="Name three plants.",
        question_ur="تین پودوں کے نام لکھیں۔",
        options_en=None,
        options_ur=None,
    )
    papers_repository.insert(
        paper_id="p1",
        subject="Science",
        class_name="Grade 4",
        total_marks=7,
        question_ids=["q1", "q2"],
    )


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_docx_unknown_paper_returns_none(test_db):
    assert export_service.build_paper_docx("nope") is None


def test_docx_is_valid_openable_document(test_db):
    _setup_paper(test_db)
    data = export_service.build_paper_docx("p1")
    assert data is not None
    # If the bytes aren't a valid docx, Document() raises — that's the assertion.
    doc = Document(io.BytesIO(data))
    assert doc is not None


def test_docx_contains_questions_marks_and_title(test_db):
    _setup_paper(test_db)
    doc = Document(io.BytesIO(export_service.build_paper_docx("p1")))
    text = _all_text(doc)
    assert "Science — Examination Paper" in text
    assert "Q1." in text and "Q2." in text
    assert "What is photosynthesis?" in text
    assert "[2]" in text and "[5]" in text
    assert "Grade 4" in text  # class_name in meta block
    assert "Total Marks" in text


def test_docx_renders_school_name_from_settings(test_db):
    _setup_paper(test_db)
    settings_repository.upsert(
        school_name="Happy Birds School",
        school_name_ur="ہیپی برڈز اسکول",
        address="Swat, KPK",
        address_ur="",
        logo_base64=None,
        accent_color="#123456",
    )
    doc = Document(io.BytesIO(export_service.build_paper_docx("p1")))
    text = _all_text(doc)
    assert "Happy Birds School" in text
    assert "ہیپی برڈز اسکول" in text


def test_docx_wires_urdu_rtl_and_complex_script_font(test_db):
    _setup_paper(test_db)
    doc = Document(io.BytesIO(export_service.build_paper_docx("p1")))
    body_xml = doc.element.body.xml
    # Urdu runs must carry the rtl flag, the bidi paragraph mark, and the
    # complex-script font — without these Word lays Urdu out left-to-right.
    assert "w:rtl" in body_xml
    assert "w:bidi" in body_xml
    assert export_service.URDU_FONT in body_xml


def test_docx_handles_question_with_no_options(test_db):
    """q2 has no options — should still render (answer-space path), not crash."""
    _setup_paper(test_db)
    doc = Document(io.BytesIO(export_service.build_paper_docx("p1")))
    assert "Name three plants." in _all_text(doc)


def test_docx_survives_corrupt_logo(test_db):
    _setup_paper(test_db)
    settings_repository.upsert(
        school_name="S",
        school_name_ur="",
        address="",
        address_ur="",
        logo_base64="data:image/png;base64,not-real-base64!!!",
        accent_color="#0e4d3c",
    )
    # Bad logo must be skipped, not fatal.
    assert export_service.build_paper_docx("p1") is not None


def test_pdf_unknown_paper_returns_none(test_db):
    assert export_service.build_paper_pdf("nope") is None


def test_pdf_raises_when_libreoffice_missing(test_db, monkeypatch):
    _setup_paper(test_db)
    monkeypatch.setattr(export_service, "_find_soffice", lambda: None)
    with pytest.raises(PdfConversionFailed, match="LibreOffice"):
        export_service.build_paper_pdf("p1")


# ---- pure helper unit tests -------------------------------------------------


def test_clean_hex_strips_hash_and_uppercases():
    assert export_service._clean_hex("#0e4d3c") == "0E4D3C"
    assert export_service._clean_hex("ABCDEF") == "ABCDEF"
    assert export_service._clean_hex("") is None
    assert export_service._clean_hex("xyz") is None  # not 6 hex chars


def test_safe_json_list_tolerates_garbage():
    assert export_service._safe_json_list('["a","b"]') == ["a", "b"]
    assert export_service._safe_json_list(None) == []
    assert export_service._safe_json_list("not json") == []
    assert export_service._safe_json_list('{"k": 1}') == []  # not a list


def test_decode_logo_handles_data_url_and_raw():
    # "AAAA" base64-decodes to 3 bytes; data-URL prefix must be stripped first.
    assert export_service._decode_logo("data:image/png;base64,AAAA") == b"\x00\x00\x00"
    assert export_service._decode_logo("AAAA") == b"\x00\x00\x00"
    assert export_service._decode_logo(None) is None
